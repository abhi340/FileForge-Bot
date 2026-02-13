from pathlib import Path
from docx import Document
from app.config import logger


class DOCXService:

    @staticmethod
    async def remove_metadata(input_path, output_path):
        doc = Document(str(input_path))
        core = doc.core_properties
        core.author = ""
        core.title = ""
        core.subject = ""
        core.keywords = ""
        core.comments = ""
        core.last_modified_by = ""
        core.category = ""
        core.content_status = ""
        doc.save(str(output_path))
        logger.info("DOCX metadata removed")
        return output_path

    @staticmethod
    async def remove_comments(input_path, output_path):
        doc = Document(str(input_path))
        body = doc.element.body
        comment_tags = ("commentRangeStart", "commentRangeEnd", "commentReference")
        to_remove = []
        for element in body.iter():
            tag_name = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag_name in comment_tags:
                to_remove.append(element)
        for element in to_remove:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        doc.save(str(output_path))
        logger.info("DOCX comments removed")
        return output_path

    @staticmethod
    async def extract_text(input_path):
        doc = Document(str(input_path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        result = "\n".join(parts)
        if not result.strip():
            result = "No extractable text found."
        logger.info(f"DOCX text extracted: {len(result)} chars")
        return result

    @staticmethod
    async def to_pdf(input_path, output_path):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors

        doc = Document(str(input_path))

        pdf_doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=20 * mm,
            leftMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        story = []

        for para in doc.paragraphs:
            if para.text.strip():
                style = styles["Normal"]
                if para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name[-1])
                        if level <= 3:
                            style = styles[f"Heading{level}"]
                    except (ValueError, KeyError):
                        style = styles["Heading1"]
                story.append(Paragraph(para.text, style))
                story.append(Spacer(1, 3 * mm))

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                max_cols = max(len(r) for r in table_data)
                for r in table_data:
                    while len(r) < max_cols:
                        r.append("")
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("PADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(t)
                story.append(Spacer(1, 5 * mm))

        if not story:
            story.append(Paragraph("Empty document", styles["Normal"]))

        pdf_doc.build(story)
        logger.info("DOCX converted to PDF")
        return output_path

    @staticmethod
    async def get_info(input_path):
        doc = Document(str(input_path))
        core = doc.core_properties

        total_text = ""
        for para in doc.paragraphs:
            total_text += para.text + " "

        words = len(total_text.split())
        chars = len(total_text)
        chars_no_space = len(total_text.replace(" ", ""))

        info = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
            "words": words,
            "characters": chars,
            "characters_no_space": chars_no_space,
            "author": core.author or "N/A",
            "title": core.title or "N/A",
            "subject": core.subject or "N/A",
            "created": str(core.created) if core.created else "N/A",
            "modified": str(core.modified) if core.modified else "N/A",
            "last_modified_by": core.last_modified_by or "N/A",
            "size_bytes": input_path.stat().st_size,
        }

        # Count images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        info["images"] = image_count

        logger.info("DOCX info retrieved")
        return info

    @staticmethod
    async def word_count(input_path):
        doc = Document(str(input_path))
        total_text = ""
        for para in doc.paragraphs:
            total_text += para.text + " "

        words = total_text.split()
        word_count = len(words)
        char_count = len(total_text)
        char_no_space = len(total_text.replace(" ", ""))
        line_count = len([p for p in doc.paragraphs if p.text.strip()])
        sentence_count = total_text.count(".") + total_text.count("!") + total_text.count("?")

        logger.info(f"Word count: {word_count}")
        return {
            "words": word_count,
            "characters": char_count,
            "characters_no_space": char_no_space,
            "lines": line_count,
            "sentences": sentence_count,
            "avg_word_length": round(char_no_space / word_count, 1) if word_count > 0 else 0,
        }

    @staticmethod
    async def extract_images(input_path, output_dir):
        output_dir.mkdir(parents=True, exist_ok=True)
        doc = Document(str(input_path))
        paths = []
        img_count = 0

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_count += 1
                image = rel.target_part
                ext = image.content_type.split("/")[-1]
                if ext == "jpeg":
                    ext = "jpg"
                img_path = output_dir / f"image_{img_count}.{ext}"
                with open(img_path, "wb") as f:
                    f.write(image.blob)
                paths.append(img_path)

        logger.info(f"DOCX images extracted: {len(paths)}")
        return paths

    @staticmethod
    async def extract_tables_csv(input_path, output_dir):
        import csv
        output_dir.mkdir(parents=True, exist_ok=True)
        doc = Document(str(input_path))
        paths = []

        for idx, table in enumerate(doc.tables):
            csv_path = output_dir / f"table_{idx + 1}.csv"
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                for row in table.rows:
                    writer.writerow([cell.text.strip() for cell in row.cells])
            paths.append(csv_path)

        logger.info(f"DOCX tables extracted: {len(paths)} CSV files")
        return paths
